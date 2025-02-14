"""
This file contains the controllers for the resume generator application
authors: Erin Hwang
"""

from app.utils.logger import LoggerConfig
from langchain.text_splitter import MarkdownHeaderTextSplitter
from docx import Document
import re
import json
import os

GENERATED_RESUME_PATH = os.getenv("GENERATED_RESUME_PATH")

#needs to read in docx file pertaining to UUID in main.py

class ResumeRendererController:
    ALLOWED_SECTIONS = {
        "core_expertise": r"(Core Expertise: )(.*)",
        "technical_snapshot": r"(Technical Snapshot: )(.*)",
        }

    pattern = r"\[[^\[\]]*\]"
    start_pattern = r"\[(?![^\[\]]*\])[^]]*$"
    end_pattern = r"^(?!.*?\[[^\[\]]*\]).*?\]"
    middle_pattern = r"^[^\[\]]*$"
    """
    Render a resume from a template and extracted content

    Args:
        resume_path (str): Path to base resume
        extracted_content (str): Generated content from ResumeGeneratorController

    Returns:
        #TODO: figure this out
    """

    def __init__(self, resume_path: str, generated_content: dict, source_name:str, md_info: str):
        self.logger = LoggerConfig().get_logger(__name__)
        self.source_name = source_name
        self.data_dir = GENERATED_RESUME_PATH
        self.resume_path = resume_path
        self.generated_content = generated_content
        self.keyword_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[("#", "keywords"),],
            strip_headers=True)

        self.cl_keyword_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[("#", "header"),],
            strip_headers=False)

        self.md_info = md_info

    def align_text(self):
        """Cleanse the markdown prefix and suffix text"""
        final = {}
        keys = ["[CompanyName]", "[CityState]", "[PositionName]"]
        cleansed_text = self.md_info.replace("```markdown", "").replace("```", "").strip()
        splits = self.cl_keyword_splitter.split_text(cleansed_text)
        assert len(splits) == 3, "Cover letter md output does not contain 3 headers"
        keywords = [each.metadata["header"] for each in splits]

        for each in list(zip(keys, keywords)):
            if "Not Available" in each[1]:
                final[each[0]] = ""
            else:
                final[each[0]] = each[1]
        return final

    def cleanse_text(self, text: str):
        """Cleanse the markdown prefix and suffix text"""
        cleansed_text = text.replace("```markdown", "").replace("```", "").strip()
        split_text = self.keyword_splitter.split_text(cleansed_text)
        return split_text[0].page_content

    def find_placement(self, doc: Document, section_name: str):
        """Find the index of the specified section in the resume"""
        curr = 0
        paragraph_found = False
        while not paragraph_found:
            if section_name in doc.paragraphs[curr].text:
                paragraph_found = True
            curr += 1
        return curr

    def render_position_sentence(self, doc: Document, curr: int, new_text: str):
        for i, run in enumerate(doc.paragraphs[curr].runs):
            if i == 0:
                run.text = new_text
            else:
                run.text = ""

    def render_keywords(self, doc: Document, section_name: str, new_text: str):
        """Render the keywords section of the resume"""
        if section_name not in self.ALLOWED_SECTIONS:
            self.logger.error(f"Invalid driver: {section_name}")
            raise ValueError(f"Invalid driver: {section_name}")

        for paragraph in doc.paragraphs:
            # create full text of the paragraph
            full_text = "".join([run.text for run in paragraph.runs])

            # retrieve exact regex pattern based on the section name aka prompt name
            pattern = self.ALLOWED_SECTIONS[section_name]
            matches = re.search(pattern, full_text)

            if matches:
                # static: prompt section/name, replaceable: base section
                static_part = matches.group(1)
                replaceable_part = matches.group(2)

                # go through each run to determine if the entirety of the static is seen
                curr_text = ""  # Track cumulative text in the paragraph
                for run in paragraph.runs:
                    # if not run.text.isspace():
                    curr_text += run.text

                    # if this logic is met, we can guarantee the next run will contain parts of the replaceable
                    if run.text.isspace() and static_part in curr_text:
                        continue
                    if replaceable_part in curr_text:
                        run.text = new_text
                        self.logger.info(f"\tUpdated {section_name} section")
                        return
        self.logger.error(f"Could not find {section_name} section in the base resume")
        raise ValueError(f"Could not find {section_name} section in the base resume")

    # TODO: find correct type for paragraph
    def return_indices_to_remove(
            self, placement: int, paragraphs, alignment_indicator: str, next_section_name: str
            ):
        """Returns paragraph indices to remove based on alignment type within the job specific section"""
        removal_indices = []
        for i in range(placement, len(paragraphs)):
            if next_section_name in paragraphs[i].text:
                return removal_indices
            if paragraphs[i].alignment.__str__() == alignment_indicator:
                removal_indices.append(i)

    # TODO: find correct type for paragraph
    def remove_paragraphs(self, paragraphs, removal_indices: list):
        for idx in reversed(removal_indices):
            p = paragraphs[idx]._element
            p.getparent().remove(p)

    #TODO: find correct type for paragraph and reference_run
    def overwrite_runs_with_text(self, paragraph, new_text:str, reference_run):
        """Overwrites bullet points from the base resume with the new generated text"""
        # Clear all existing runs
        for r in paragraph.runs:
            r.text = ""

        # Add one run with the new text
        new_run = paragraph.add_run(new_text)
        new_run.font.bold = reference_run.font.bold
        new_run.font.italic = reference_run.font.italic
        new_run.font.underline = reference_run.font.underline
        new_run.font.size = reference_run.font.size
        if reference_run.font.color is not None:
            new_run.font.color.rgb = reference_run.font.color.rgb
        new_run.font.name = reference_run.font.name

    def decouple_professional_experience(self, generated_str: str):
        """Decouples the high level sentence from the bullet points"""
        cleansed_professional = [
            each.removeprefix("- ").removeprefix("* ") for each in generated_str.split('\n')
            ]
        return cleansed_professional[0], cleansed_professional[1:]

    def cleanse_generated_content(self):
        """Cleanse the generated content for rendering"""
        cleansed_content = {}
        for section_name, generated_content in self.generated_content.items():
            if ' ' in section_name:
                #this is the professional experience section
                cleansed_professional_experience = self.cleanse_text(generated_content)
                high_level_sentence, bullet_points = self.decouple_professional_experience(
                    cleansed_professional_experience
                    )
                cleansed_content[section_name] = {
                    "upper": high_level_sentence, "lower": bullet_points
                    }

            else:
                #this is the core expertise and technical snapshot section
                cleansed_content[section_name] = self.cleanse_text(generated_content)

        self.logger.info("Generated content has been cleansed - ready for rendering")
        return cleansed_content

    #TODO: find correct type for paragraph
    def overwrite_entire_section(self, keep_indices: list, paragraphs, list_bullets: list[str]):
        """Overwrites the entire section with the new generated content"""
        for idx_par, text_str in zip(keep_indices, list_bullets):
            kept_paragraph = paragraphs[idx_par]

            alignment_type = kept_paragraph.alignment #TODO

            if not alignment_type:
                # if no runs exist, create a dummy run to copy formatting from
                if not kept_paragraph.runs:
                    reference_run = kept_paragraph.add_run("")
                else:
                    reference_run = kept_paragraph.runs[0]

                # overwrite that paragraph's text with text_str
                self.overwrite_runs_with_text(kept_paragraph, text_str, reference_run)
            else:
                self.logger.info(f"\tSkipping paragraph at index {idx_par}; alignment is None. Not a bullet paragraph.")

    def render_professional_experience(
            self, doc: Document, current_section_name: str, next_section_name: str,
            upper: str, lower: list
            ):
        """Renders a single professional experience using the LLM-generated content"""
        #1: high level sentence rendering
        upper_placement = self.find_placement(doc, current_section_name)
        self.render_position_sentence(doc, upper_placement, upper)
        self.logger.info(f"\tRendered high level sentence for position title: {current_section_name}")

        #2: bullet point rendering
        all_paragraphs = doc.paragraphs
        x = len(lower) #number of bullet points paragraphs to overwrite

        #2.1 find the indices to remove (strict) then execute removal
        strict_removal_indices = self.return_indices_to_remove(
            upper_placement + 1, all_paragraphs, "None", next_section_name
            )
        self.remove_paragraphs(all_paragraphs, strict_removal_indices)
        all_paragraphs = doc.paragraphs

        #2.2find indices to remove that are bullet points then only keep the first x
        removal_indices = self.return_indices_to_remove(
            upper_placement + 1, all_paragraphs, "LEFT (0)", next_section_name
            )
        keep_count = min(x, len(removal_indices))
        indices_to_keep = removal_indices[:keep_count] #keep the first x from starting upper placement
        indices_to_remove = [idx for idx in removal_indices[keep_count:]] #delete the rest up until the next section

        self.remove_paragraphs(all_paragraphs, indices_to_remove)
        all_paragraphs = doc.paragraphs
        self.logger.info(
            "\tSuccessfully removed %s bullet points for position title: %s",
            len(indices_to_remove),
            current_section_name
            )

        #2.3 now overwrite the runs that were left out from above
        self.overwrite_entire_section(indices_to_keep, all_paragraphs, lower)
        self.logger.info(
            "\tRendered %s bullet points for position title: %s", str(x), current_section_name
            )

    def edit_paragraphs(self, all_paragraphs, i_edit_list: list[int], content_dict: dict):
        cleansed_content = content_dict
        for i_par_edit in i_edit_list:
            # print(i_par_edit)
            par = all_paragraphs[i_par_edit]
            # print(f"Paragraph {i_par} text:\n{par.text}\n-Here are the runs:")

            full_matches = []
            start_matches = []
            end_matches = []
            middle_matches = []



            for i_run, run in enumerate(par.runs):
                # print(run.text)
                full_match = [(i_run, val) for val in re.findall(self.pattern, run.text)]
                full_matches.extend(full_match)
                start_match =  [(i_run, val) for val in re.findall(self.start_pattern, run.text)]
                start_matches.extend(start_match)
                end_match =  [(i_run, val) for val in re.findall(self.end_pattern, run.text)]
                end_matches.extend(end_match)
                if full_match or start_match:
                    middle_match = [(i_run, val) for val in re.findall(self.middle_pattern, run.text)]
                    middle_matches.extend(middle_match)

            if full_matches:
                # print("here")
                # print(full_matches)
                for each_full_match in full_matches:
                    # print(each_full_match[0])
                    all_paragraphs[i_par_edit].runs[each_full_match[0]].text = all_paragraphs[i_par_edit].runs[each_full_match[0]].text.replace(each_full_match[1], cleansed_content[each_full_match[1]])
                    all_paragraphs[i_par_edit].runs[each_full_match[0]].font.name = "Calibri"
                    all_paragraphs[i_par_edit].runs[each_full_match[0]].font.size = 133350

            if start_matches:
                assert len(start_matches) == len(end_matches)

                i_curr = 0
                for start, end in zip(start_matches, end_matches):

                    i_start = start[0]
                    val_start = start[1]

                    i_end = end[0]
                    val_end = end[1]

                    curr_middle_match = None if len(middle_matches) == 0 else middle_matches[i_curr]

                    if i_end - i_start >= 2:
                        assert curr_middle_match is not None
                        determined_term = val_start + curr_middle_match[1] + val_end
                        all_paragraphs[i_par_edit].runs[curr_middle_match[0]].text = all_paragraphs[i_par_edit].runs[curr_middle_match[0]].text.replace(curr_middle_match[1], "")

                        all_paragraphs[i_par_edit].runs[curr_middle_match[0]].font.name = "Calibri"
                        all_paragraphs[i_par_edit].runs[curr_middle_match[0]].font.size = 133350
                    else:
                        assert curr_middle_match is None
                        determined_term = val_start + val_end

                    # for
                    i_curr += 1

                    #replacing at the run where [ is first found
                    all_paragraphs[i_par_edit].runs[i_start].text = all_paragraphs[i_par_edit].runs[i_start].text.replace(val_start, cleansed_content[determined_term])
                    # all_paragraphs[i_par_edit].runs[each_full_match[0]].font.name = "Calibri"
                    # all_paragraphs[i_par_edit].runs[each_full_match[0]].font.size = 133350

                    #replace at the run where ] is found
                    all_paragraphs[i_par_edit].runs[i_end].text = all_paragraphs[i_par_edit].runs[i_end].text.replace(val_end, "")
                    # all_paragraphs[i_par_edit].runs[each_full_match[0]].font.name = "Calibri"
                    # all_paragraphs[i_par_edit].runs[each_full_match[0]].font.size = 133350

    def render_summary(self, doc: Document, summary_str: str):
        #find the run that contains the text
        for i_par, paragraph in enumerate(doc.paragraphs):
            for i_run, run in enumerate(paragraph.runs):
                if run.text.__contains__("As a multitalented Data Scientist"):
                    #overwrite the run text with summary str
                    doc.paragraphs[i_par].runs[i_run].text = summary_str
                    return


    @LoggerConfig().log_execution
    def execute(self):
        """Execute resume rendering process"""
        cleansed_cl_keywords = self.align_text()

        doc = Document(self.resume_path)
        #replace [Keywords] with the actual keyword
        i_pars_to_edit = [i_par for i_par, par in enumerate(doc.paragraphs) if par.text.__contains__("[")]
        self.edit_paragraphs(doc.paragraphs, i_pars_to_edit, cleansed_cl_keywords)
        self.logger.info("Keywords have been rendered")

        cleansed_content = self.cleanse_generated_content()
        previous_titles = [name for name in list(cleansed_content) if ' ' in name]
        #add render professional summary here
        self.render_summary(doc, cleansed_content["professional_summary"])
        self.render_keywords(doc, "core_expertise", cleansed_content["core_expertise"])
        self.render_keywords(doc, "technical_snapshot", cleansed_content["technical_snapshot"])

        self.logger.info("Rendering the professional experience...")

        i_curr = 0
        while i_curr < len(previous_titles):
            i_next = i_curr + 1
            #retrieving section names to set up renderer
            curr_section = previous_titles[i_curr]
            #TODO: high priority to edit if other people want to use this app
            next_section = previous_titles[i_next] if i_next != len(previous_titles) else "Education"

            #retrieve the generated high-level description of job title A
            upper = cleansed_content[curr_section]["upper"]
            #retrieve the generated bullet points of job title A
            lower = cleansed_content[curr_section]["lower"]
            self.render_professional_experience(doc, curr_section, next_section, upper, lower)
            i_curr += 1

        # self.render_professional_experience(doc, curr_section, next_section, upper, lower)
        docx_fp = f"{self.data_dir}/{self.source_name}.docx" #ERIN
        docx_fp = os.path.abspath(docx_fp)
        doc.save(docx_fp)
        self.logger.info(f"Resume rendered successfully at: {docx_fp}")
        return f"{self.data_dir}/{self.source_name}.docx"



if __name__ == "__main__":
    generated_content_fp = "/Users/erinhwang/Projects/ResuMate/experiments/generator_content.json"
    with open(generated_content_fp, "r") as f:
        content = json.load(f)

    resume_fp = "/Users/erinhwang/Projects/ResuMate/experiments/thee_resume_rendrr.docx"
    renderer = ResumeRendererController(resume_fp, content)
    renderer.execute()
